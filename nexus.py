import os
import time
import torch
import re
import pandas as pd
from typing import Tuple, List, Dict, Optional
import fitz
from docx import Document as DocxDocument
from llama_index.core import Document, Settings, VectorStoreIndex
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from threading import Lock
import hashlib
import shutil
import logging

# Suppress verbose warnings
logging.getLogger("transformers").setLevel(logging.ERROR)

from llama_index.core import StorageContext, load_index_from_storage
from llama_index.llms.huggingface import HuggingFaceLLM
from llama_index.embeddings.langchain import LangchainEmbedding
from langchain_huggingface import HuggingFaceEmbeddings
from transformers import AutoTokenizer, AutoModelForCausalLM
from rich.console import Console
from rich.panel import Panel
from rich.rule import Rule
from rich.markdown import Markdown
from rich.prompt import Prompt

# --- Globals ---
PERSIST_DIR = "index_storage"
FILES_DIR = "files"
console = Console()
documents = []
query_engine = None
file_snapshot = {}
lock = Lock()


def dataframe_to_markdown(df: pd.DataFrame, filename: str) -> str:
    """Convert a pandas DataFrame to a Markdown table."""
    MAX_ROWS = 50
    preview_df = df.head(MAX_ROWS)
    metadata = [
        f"### Table Preview from `{filename}`",
        f"- **Shape**: {df.shape[0]} rows × {df.shape[1]} columns",
        "- **Columns**:",
        *(f"  - `{col}`: {dtype}" for col, dtype in zip(df.columns, df.dtypes)),
        "",
        preview_df.to_markdown(index=False)
    ]
    return "\n".join(metadata)

# --- File loaders ---
FILE_LOADERS = {
    '.txt': lambda p: open(p, 'r', encoding='utf-8').read(),
    '.pdf': lambda p: "".join(page.get_text() for page in fitz.open(p)),
    '.docx': lambda p: "\n".join(para.text for para in DocxDocument(p).paragraphs),
    '.csv': lambda p: dataframe_to_markdown(pd.read_csv(p), os.path.basename(p)),
    '.xls': lambda p: dataframe_to_markdown(pd.read_excel(p), os.path.basename(p)),
    '.xlsx': lambda p: dataframe_to_markdown(pd.read_excel(p), os.path.basename(p)),
}

# --- Embedding model ---
def get_embed_model():
    return LangchainEmbedding(
        HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    )

# --- LLM factory with fallback and no quantization ---
def create_llm(model_name: str) -> HuggingFaceLLM:
    use_cuda = torch.cuda.is_available()
    dtype = torch.float16 if use_cuda else torch.float32

    tokenizer = AutoTokenizer.from_pretrained(model_name)
    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token

    model = AutoModelForCausalLM.from_pretrained(
        model_name,
        device_map="auto" if use_cuda else None,
        dtype=dtype,
    )
    if not use_cuda:
        model = model.to("cpu")
    if model.config.pad_token_id is None:
        model.config.pad_token_id = tokenizer.pad_token_id
    if model.config.eos_token_id is None:
        model.config.eos_token_id = tokenizer.eos_token_id

    return HuggingFaceLLM(
        model=model,
        tokenizer=tokenizer,
        max_new_tokens=300,
        context_window=2048,
        device_map="auto" if use_cuda else "cpu",
        system_prompt="You are Nexus, a helpful document assistant. Answer questions based only on the provided documents. Be concise and direct.",
        query_wrapper_prompt="Question: {query_str} Answer: ",
    )

def get_llm() -> HuggingFaceLLM:
    candidates = [
        "microsoft/phi-2",
        "microsoft/phi-1_5",
        "microsoft/phi-1"
    ]
    for name in candidates:
        try:
            console.print(f"[dim]Trying {name}...[/dim]")
            llm = create_llm(name)
            test = llm.complete("What is 2+2?")
            if test.text and test.text.strip():
                console.print(f"[green]✔ Using {name} – test: 2+2 = {test.text.strip()}[/green]")
                return llm
            else:
                console.print(f"[yellow]⚠ {name} returned empty, trying next...[/yellow]")
        except Exception as e:
            console.print(f"[yellow]⚠ {name} failed: {e}, trying next...[/yellow]")
    console.print("[bold red]❌ No working Phi model found. Please check your installation.")
    exit(1)

# --- Chunking (smaller to avoid token overflow) ---
from llama_index.core.node_parser import SimpleNodeParser

def chunk_documents(documents: List[Document]) -> List[Document]:
    parser = SimpleNodeParser.from_defaults(chunk_size=128, chunk_overlap=20)
    nodes = []
    for doc in documents:
        try:
            nodes.extend(parser.get_nodes_from_documents([doc]))
        except Exception as e:
            console.print(f"[bold red]Failed to chunk {doc.metadata.get('file_name', '')}: {e}[/bold red]")
    return nodes

def load_documents(directory: str) -> List[Document]:
    docs = []
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext in FILE_LOADERS:
            try:
                content = FILE_LOADERS[ext](file_path)
                docs.append(Document(text=content, metadata={"file_name": filename}))
                console.print(f"[green]Loaded: {filename}[/green]")
            except Exception as e:
                console.print(f"[bold red]Failed to load {filename}: {e}[/bold red]")
        else:
            console.print(f"[yellow]Skipped unsupported: {filename}[/yellow]")
    return docs

# --- Indexing with tree_summarize ---
def build_index_and_engine(documents: List[Document], persist: bool = True):
    nodes = chunk_documents(documents)
    storage_context = StorageContext.from_defaults()
    index = VectorStoreIndex.from_documents([], storage_context=storage_context)
    index.insert_nodes(nodes)
    if persist:
        index.storage_context.persist(persist_dir=PERSIST_DIR)
    return index, index.as_query_engine(response_mode="tree_summarize", verbose=True)

def load_or_build_index(documents: List[Document]):
    required_files = ["docstore.json", "vector_store.json", "index_store.json"]
    if all(os.path.exists(os.path.join(PERSIST_DIR, f)) for f in required_files):
        try:
            console.print("[bold]Loading existing index from disk...[/bold]")
            storage_context = StorageContext.from_defaults(persist_dir=PERSIST_DIR)
            index = load_index_from_storage(storage_context)
            return index, index.as_query_engine(response_mode="tree_summarize", verbose=True)
        except Exception as e:
            console.print(f"[bold yellow]Failed to load index: {e}. Rebuilding...[/bold yellow]")
    console.print("[bold]Building new index...[/bold]")
    return build_index_and_engine(documents)

# --- File monitoring helpers ---
def hash_file(filepath: str) -> str:
    hasher = hashlib.sha256()
    with open(filepath, 'rb') as f:
        while chunk := f.read(8192):
            hasher.update(chunk)
    return hasher.hexdigest()

def get_files_snapshot(directory: str) -> Dict[str, str]:
    snapshot = {}
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            try:
                snapshot[filename] = hash_file(file_path)
            except Exception as e:
                console.print(f"[bold red]Failed to hash {filename}: {e}[/bold red]")
    return snapshot

def parse_query(user_query: str) -> Tuple[Optional[str], str]:
    file_match = re.search(
        r"(?:from|in|of|regarding|about|open|for|in the document|in file)\s+[\"']?([\w\-\s]+\.(?:txt|pdf|docx|csv|xlsx?|xls))[\"']?",
        user_query, re.IGNORECASE
    )
    if file_match:
        file_name = file_match.group(1).strip()
        stripped_query = user_query.replace(file_match.group(0), "").strip()
        return file_name, stripped_query or f"Summarize the contents of {file_name}"
    trailing_match = re.search(r"([\w\-\s]+\.(?:txt|pdf|docx|csv|xlsx?|xls))$", user_query.strip(), re.IGNORECASE)
    if trailing_match:
        file_name = trailing_match.group(1).strip()
        return file_name, f"Summarize the contents of {file_name}"
    return None, user_query

class FileChangeHandler(FileSystemEventHandler):
    def on_any_event(self, event):
        global documents, query_engine, file_snapshot
        if event.is_directory:
            return
        with lock:
            current_snapshot = get_files_snapshot(FILES_DIR)
            if current_snapshot != file_snapshot:
                console.print("\n[bold yellow]File change detected. Reloading...[/bold yellow]")
                with console.status("[bold green]Processing files...", spinner="dots"):
                    documents = load_documents(FILES_DIR)
                    if documents:
                        _, query_engine = load_or_build_index(documents)
                        file_snapshot = current_snapshot
                console.print("[bold green]Reload complete.[/bold green]")

def normalize(name: str) -> str:
    return name.strip().lower().replace(" ", "")

def show_intro():
    os.system('cls' if os.name == 'nt' else 'clear')
    art = r"""
███╗   ██╗███████╗██╗  ██╗██╗   ██╗███████╗
████╗  ██║██╔════╝╚██╗██╔╝██║   ██║██╔════╝
██╔██╗ ██║█████╗   ╚███╔╝ ██║   ██║███████╗
██║╚██╗██║██╔══╝   ██╔██╗ ██║   ██║╚════██║
██║ ╚████║███████╗██╔╝ ██╗╚██████╔╝███████║
╚═╝  ╚═══╝╚══════╝╚═╝  ╚═╝ ╚═════╝ ╚══════╝
    """
    console.print(art, style="bold blue")
    welcome_message = (
        "Welcome to [bold][cyan]Nexus[/cyan][/bold] – your intelligent document companion, powered by [bold][cyan]Microsoft Phi[/cyan][/bold].\n"
        "Developed by [bold green]Soumyadeep Basu[/bold green]\n\n"
        "[italic red]This Project is a part of my Undergrad Engineering Curriculum.[/italic red]"
        "\n\n[bold]Instructions:[/bold]\n"
        "1. Place your documents[red] BEFORE [/red]in the 'files/' directory.\n"
        "2. Ask questions using natural language.\n"
        "3. Use [bold]\"info\"[/bold] to see project details.\n"
        "4. Use [bold]\"cls\"[/bold] to clear the screen.\n"
        "5. Use [bold]\"exit\"[/bold] or [bold]\"quit\"[/bold] to end the session."
    )
    console.print(Panel(welcome_message, title="[cyan]Introduction[/cyan]", border_style="cyan", title_align="center"))

def show_info():
    info_message = (
        "[cyan]Nexus[/cyan] is a local, document‑aware chatbot that helps you retrieve insights from your files.\n"
        "It runs entirely on your machine using [bold][cyan]Microsoft Phi[/cyan][/bold] – no cloud, no tokens.\n\n"
        "Developed by [bold green]Soumyadeep Basu[/bold green].\n\n"
        "Under the hood:\n"
        "- [yellow]Chunking[/yellow]: splits documents into 128‑char pieces with 20‑char overlap.\n"
        "- [yellow]Embeddings[/yellow]: uses 'all-mpnet-base-v2' to create semantic vectors.\n"
        "- [yellow]Retrieval[/yellow]: finds the most relevant chunks for your query.\n"
        "- [yellow]Generation[/yellow]: Phi synthesises an answer from those chunks.\n\n"
        "All processing stays on your computer – your data never leaves."
    )
    console.print(Panel(info_message, title="[cyan]Project Information[/cyan]", border_style="cyan", title_align="center"))

# --- Main ---
def main():
    global documents, query_engine, file_snapshot

    # Cleanup old index
    if os.path.exists(PERSIST_DIR):
        console.print(f"[bold yellow]Removing old index...[/bold yellow]")
        try:
            shutil.rmtree(PERSIST_DIR)
        except Exception as e:
            console.print(f"[bold red]Error removing {PERSIST_DIR}: {e}")
            return

    if not os.path.exists(FILES_DIR):
        os.makedirs(FILES_DIR)
        console.print(f"[bold yellow]Created '{FILES_DIR}'. Add documents and restart.")
        return

    show_intro()

    with console.status("[bold green]Loading models...", spinner="dots"):
        llm = get_llm()
        Settings.llm = llm
        Settings.embed_model = get_embed_model()
        Settings.chunk_size = 128
        Settings.context_window = 2048

    with console.status("[bold green]Scanning and loading documents...", spinner="dots"):
        documents = load_documents(FILES_DIR)

    if not documents:
        console.print("[bold red]No supported documents found. Add files to 'files/' and restart.")
        return

    with console.status("[bold green]Building index...", spinner="dots"):
        _, query_engine = load_or_build_index(documents)
        file_snapshot = get_files_snapshot(FILES_DIR)

    console.print("\n[bold green]✅ Ready! Ask away.[/bold green]")

    event_handler = FileChangeHandler()
    observer = Observer()
    observer.schedule(event_handler, path=FILES_DIR, recursive=False)
    observer.start()

    try:
        while True:
            console.print(Rule(style="dim white"))
            user_query = Prompt.ask("[bold green]>").strip()
            if user_query.lower() in ['exit', 'quit']:
                console.print("[bold yellow]Goodbye![/bold yellow]")
                break
            if user_query.lower() == 'cls':
                show_intro()
                continue
            if user_query.lower() == 'info':
                show_info()
                continue

            response_str = ""
            with lock:
                with console.status("[bold cyan]Thinking...", spinner="earth"):
                    target_file, stripped_query = parse_query(user_query)
                    if target_file:
                        matching_docs = [doc for doc in documents if normalize(doc.metadata.get("file_name", "")) == normalize(target_file)]
                        if matching_docs:
                            _, filtered_engine = build_index_and_engine(matching_docs, persist=False)
                            response = filtered_engine.query(stripped_query)
                            response_str = str(response).strip() if response else ""
                        else:
                            response_str = f"No document named '{target_file}' found."
                    else:
                        response = query_engine.query(user_query)
                        response_str = str(response).strip() if response else ""

            if not response_str:
                response_str = "I couldn't find an answer. Try rephrasing or asking a more specific question."

            console.print(Panel(Markdown(response_str), border_style="blue", title="[bold cyan]Nexus[/bold cyan]", title_align="left"))

    finally:
        observer.stop()
        observer.join()
        if os.path.exists(PERSIST_DIR):
            shutil.rmtree(PERSIST_DIR)
            console.print("[green]Cleaned up index.[/green]")
        console.print("[bold yellow]Session ended.[/bold yellow]")

if __name__ == "__main__":
    main()